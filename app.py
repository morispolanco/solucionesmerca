import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Diccionario de Problemas de Mercadeo", page_icon="📚", layout="wide")

# Function to set the background color
def set_background_color(color):
    st.markdown(
        f"""
        <style>
        .stApp {{
            background-color: {color};
        }}
        </style>
        """,
        unsafe_allow_html=True
    )

# Function to create the information column
def crear_columna_info():
    st.markdown("""
    ### Sobre esta aplicación

    Esta aplicación es un Diccionario de Problemas de Mercadeo. Permite a los usuarios obtener soluciones creativas a problemas de mercadeo para diversos servicios o industrias.

    ### Cómo usar la aplicación:

    1. Elija un problema de mercadeo de la lista predefinida o proponga su propio problema.
    2. Seleccione uno o más servicios o industrias.
    3. Haga clic en "Obtener solución" para generar las respuestas.
    4. Lea las soluciones y fuentes proporcionadas.
    5. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **Moris Polanco**, 26 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Diccionario de Problemas de Mercadeo* [Aplicación web]. https://solucionesmerca.streamlit.app

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar respuestas basadas en información disponible en línea. Siempre verifique la información con fuentes académicas para un análisis más profundo.
    """)

# Titles and Main Column
st.title("Diccionario de Problemas de Mercadeo")

# Set background color to light yellow
set_background_color("#FFF9C4")  # Light yellow color code

col1, col2 = st.columns([1, 2])

with col1:
    crear_columna_info()

with col2:
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPER_API_KEY = st.secrets["SERPER_API_KEY"]

    # List of 101 marketing problems
    problemas_mercadeo = sorted([
        "¿Cómo aumentar la lealtad del cliente?", "¿Cómo mejorar el posicionamiento de marca?", "¿Cómo incrementar las ventas online?", 
        "¿Cómo mejorar el ROI de campañas de marketing?", "¿Cómo gestionar la reputación online?", "¿Cómo atraer tráfico a una tienda física?", 
        "¿Cómo optimizar un sitio web para SEO?", "¿Cómo realizar un lanzamiento de producto exitoso?", 
        "¿Cómo fidelizar a los usuarios de una aplicación móvil?", "¿Cómo crear contenido de marketing viral?", 
        "¿Cómo integrar campañas de marketing online y offline?", "¿Cómo reducir el costo por adquisición de cliente?", 
        "¿Cómo segmentar a la audiencia de manera efectiva?", "¿Cómo utilizar el marketing de influencers?", 
        "¿Cómo generar leads de calidad?", "¿Cómo utilizar el email marketing de forma efectiva?", "¿Cómo optimizar el embudo de ventas?", 
        "¿Cómo incrementar la tasa de conversión en un ecommerce?", "¿Cómo mejorar la presencia en redes sociales?", 
        "¿Cómo utilizar el video marketing?", "¿Cómo aprovechar el marketing de guerrilla?", "¿Cómo mejorar la experiencia del cliente?", 
        "¿Cómo medir la efectividad de una campaña de marketing?", "¿Cómo definir una estrategia de marketing digital?", 
        "¿Cómo realizar un análisis de competidores?", "¿Cómo gestionar una crisis de relaciones públicas?", 
        "¿Cómo utilizar el marketing de contenidos?", "¿Cómo implementar técnicas de remarketing?", 
        "¿Cómo aprovechar las tendencias de consumo?", "¿Cómo incrementar la notoriedad de marca?", 
        "¿Cómo realizar un rebranding exitoso?", "¿Cómo utilizar el marketing basado en datos?", 
        "¿Cómo mejorar la estrategia de precios?", "¿Cómo crear una campaña publicitaria efectiva?", 
        "¿Cómo manejar el marketing multicultural?", "¿Cómo utilizar el marketing de proximidad?", 
        "¿Cómo incrementar la visibilidad en motores de búsqueda?", "¿Cómo implementar estrategias de co-marketing?", 
        "¿Cómo manejar la personalización en marketing?", "¿Cómo desarrollar una propuesta de valor única?", 
        "¿Cómo aumentar las suscripciones a newsletters?", "¿Cómo utilizar la automatización de marketing?", 
        "¿Cómo optimizar el marketing en eventos?", "¿Cómo fomentar el branding emocional?", 
        "¿Cómo llevar a cabo una estrategia de marketing local?", "¿Cómo aprovechar el marketing estacional?", 
        "¿Cómo manejar las críticas y reseñas negativas?", "¿Cómo maximizar el uso de testimonios y reseñas?", 
        "¿Cómo utilizar el marketing con causa?", "¿Cómo implementar una estrategia omnicanal?", 
        "¿Cómo destacarse en un mercado saturado?", "¿Cómo utilizar la realidad aumentada en marketing?", 
        "¿Cómo mejorar la interacción en redes sociales?", "¿Cómo implementar el marketing basado en cuentas (ABM)?", 
        "¿Cómo realizar una estrategia de growth hacking?", "¿Cómo diversificar las estrategias de adquisición de clientes?", 
        "¿Cómo optimizar el uso de hashtags?", "¿Cómo manejar el marketing móvil?", "¿Cómo realizar campañas de marketing colaborativo?", 
        "¿Cómo gestionar el contenido generado por el usuario?", "¿Cómo utilizar el marketing sensorial?", 
        "¿Cómo identificar y aprovechar nichos de mercado?", "¿Cómo mejorar la satisfacción del cliente?", 
        "¿Cómo utilizar el neuromarketing?", "¿Cómo realizar campañas de marketing en tiempo real?", 
        "¿Cómo mejorar la tasa de retención de clientes?", "¿Cómo utilizar el storytelling?", 
        "¿Cómo implementar el marketing predictivo?", "¿Cómo mejorar la usabilidad de un sitio web?", 
        "¿Cómo realizar campañas de marketing social?", "¿Cómo manejar el marketing de afiliados?", 
        "¿Cómo generar engagement en redes sociales?", "¿Cómo utilizar el marketing de experiencias?", 
        "¿Cómo sacarle partido al marketing B2B?", "¿Cómo realizar una investigación de mercado efectiva?", 
        "¿Cómo lanzar una campaña de crowfunding?", "¿Cómo optimizar el marketing de boca en boca?", 
        "¿Cómo utilizar estrategias de marketing sostenible?", "¿Cómo manejar la publicidad programática?", 
        "¿Cómo gestionar una estrategia de branding digital?", "¿Cómo utilizar el marketing de clubes de fidelidad?", 
        "¿Cómo realizar campañas de marketing transmedia?", "¿Cómo implementar el marketing inclusivo?", 
        "¿Cómo utilizar el marketing de asociaciones?", "¿Cómo optimizar el marketing de productos exclusivos?", 
        "¿Cómo manejar el marketing de entrada (inbound marketing)?", "¿Cómo realizar una estrategia de marketing para startups?", 
        "¿Cómo mejorar la respuesta de campañas en distintos dispositivos?", "¿Cómo realizar campañas de marketing emocional?", 
        "¿Cómo manejar una estrategia de marketing en plataformas de streaming?", "¿Cómo utilizar el remarketing dinámico?", 
        "¿Cómo optimizar el marketing en ferias y exposiciones?", "¿Cómo manejar el marketing de adopción tecnológica?", 
        "¿Cómo utilizar estrategias de marketing en festivales?", "¿Cómo optimizar el marketing de aplicaciones SaaS?", 
        "¿Cómo optimizar la velocidad de carga de un sitio web?", "¿Cómo realizar campañas de marketing en deportes?", 
        "¿Cómo manejar una estrategia de marketing en la industria de la moda?", "¿Cómo manejar una estrategia de marketing en el sector inmobiliario?"
    ])

    # List of services or industries
    servicios_industrias = [
        "Tecnología", "Salud", "Educación", "Finanzas", "Alimentación", 
        "Moda", "Automotriz", "Turismo", "Bienes Raíces", "Deportes", 
        "Banca", "Comunicaciones", "Energía", "Transporte", "Seguros"
    ]

    def buscar_informacion(query, industria):
        url = "https://google.serper.dev/search"
        payload = json.dumps({
            "q": f"{query} {industria} mercadeo"
        })
        headers = {
            'X-API-KEY': SERPER_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()

    def generar_respuesta(problema, industria, contexto):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Contexto: {contexto}\n\nProblema: {problema}\nIndustria: {industria}\n\nProporciona una solución creativa al problema de mercadeo '{problema}' según la industria de {industria}. La solución debe ser detallada y práctica, similar a una estrategia de mercadeo innovadora. Si es posible, incluye una referencia a una técnica, estudio de caso o estrategia de {industria} que trate este concepto.\n\nSolución:",
            "max_tokens": 2048,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["Problema:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(problema, respuestas, fuentes):
        doc = Document()
        doc.add_heading('Diccionario de Problemas de Mercadeo', 0)

        doc.add_heading('Problema', level=1)
        doc.add_paragraph(problema)

        for industria, respuesta in respuestas.items():
            doc.add_heading(f'Solución para la industria {industria}', level=2)
            doc.add_paragraph(respuesta)

        doc.add_heading('Fuentes', level=1)

        # Limitar la lista de fuentes a las primeras 10
        for fuente in fuentes[:10]:
            doc.add_paragraph(fuente, style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

        return doc

    st.write("**Elige un problema de mercadeo de la lista o propón tu propio problema**:")

    opcion = st.radio("", ["Elegir de la lista", "Proponer mi propio problema"])

    if opcion == "Elegir de la lista":
        problema = st.selectbox("Selecciona un problema:", problemas_mercadeo)
    else:
        problema = st.text_input("Ingresa tu propio problema de mercadeo:")

    st.write("Selecciona uno o más servicios o industrias (máximo 5):")
    industrias_seleccionadas = st.multiselect("Servicios o Industrias", servicios_industrias)

    if len(industrias_seleccionadas) > 5:
        st.warning("Has seleccionado más de 5 servicios o industrias. Por favor, selecciona un máximo de 5.")
    else:
        if st.button("Obtener solución"):
            if problema and industrias_seleccionadas:
                with st.spinner("Buscando información y generando soluciones..."):
                    respuestas, todas_fuentes = {}, []

                    for industria in industrias_seleccionadas:
                        # Buscar información relevante
                        resultados_busqueda = buscar_informacion(problema, industria)
                        contexto = "\n".join([item["snippet"] for item in resultados_busqueda.get("organic", [])])
                        fuentes = [item["link"] for item in resultados_busqueda.get("organic", [])]

                        # Generar respuesta
                        respuesta = generar_respuesta(problema, industria, contexto)

                        respuestas[industria] = respuesta
                        todas_fuentes.extend(fuentes)

                    # Mostrar las respuestas
                    st.subheader(f"Soluciones para el problema: {problema}")
                    for industria, respuesta in respuestas.items():
                        st.markdown(f"**{industria}:** {respuesta}")

                    # Botón para descargar el documento
                    doc = create_docx(problema, respuestas, todas_fuentes)
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    st.download_button(
                        label="Descargar solución en DOCX",
                        data=buffer,
                        file_name=f"Solución_{problema.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("Por favor, selecciona un problema y al menos un servicio o industria.")
